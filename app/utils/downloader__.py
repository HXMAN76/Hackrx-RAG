from email import utils
from llama_cloud_services import LlamaParse
import asyncio
import time
import os
from dotenv import load_dotenv
import aiohttp
from pathlib import Path
import tempfile
import logging
import mimetypes
import re 
from collections import Counter
import docx
import email
from email import policy
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


load_dotenv()
from app.config import (
    LLAMA_API_KEY,
    LLAMA_LANGUAGE,
    LLAMA_FAST_MODE,
    LLAMA_DISABLE_OCR,
    LLAMA_DISABLE_IMG,
    LLAMA_HIDE_HEADERS,
    LLAMA_HIDE_FOOTERS,
    ASYNC_TIMEOUT
)

parser = LlamaParse(
    api_key=LLAMA_API_KEY,
    verbose=True,
    language=LLAMA_LANGUAGE,
    disable_ocr=LLAMA_DISABLE_OCR,
    disable_image_extraction=LLAMA_DISABLE_IMG,
    hide_headers=LLAMA_HIDE_HEADERS,
    hide_footers=LLAMA_HIDE_FOOTERS,
    fast_mode=LLAMA_FAST_MODE,
)

class UniversalTextCleaner:
    def __init__(self):
        self.noise_patterns = [
            r'UIN[:\-\s]*[A-Z0-9]+',
            r'Reg\.?\s*No\.?\s*[:\-\s]*\d+',
            r'CIN\s*[:\-\s]*[A-Z0-9]+',
            r'IRDAI\s*Regn?\.?\s*No\.?\s*[:\-\s]*\d+',
            r'Page\s+\d+\s+of\s+\d+',
            r'^\d+\s*$',
            r'www\.[a-zA-Z0-9\-\.]+\.com',
            r'E-mail:\s*[a-zA-Z0-9\-\.@]+',
            r'Call at:\s*.*?\(Toll Free.*?\)',
            r'For more details.*?Toll Free.*?\)',
            r'Regd\.?\s*&\s*Head Office:.*?-\s*\d{6}',
            r'Plot no\..*?-\s*\d{6}',
            r'^[A-Z\s]{10,}\s*$',
        ]
        self.ocr_corrections = {
            ' . ': '. ',
            ' , ': ', ',
            ' ; ': '; ',
            ' : ': ': ',
            '( ': '(',
            ' )': ')',
            ' / ': '/',
            'O ': '0 ',
            'l ': '1 ',
            '  ': ' ',
        }

    def detect_repeated_elements(self, lines, threshold=3):
        counts = Counter(line.strip() for line in lines if len(line.strip()) > 10)
        return {line for line, count in counts.items() if count >= threshold}

    def remove_noise_patterns(self, text):
        lines = text.split("\n")
        repeated = self.detect_repeated_elements(lines)
        cleaned = []
        for line in lines:
            line = line.strip()
            if not line or line in repeated:
                continue
            if any(re.search(p, line, re.IGNORECASE) for p in self.noise_patterns):
                continue
            cleaned.append(line)
        return " ".join(cleaned)

    def correct_ocr_errors(self, text):
        for err, corr in self.ocr_corrections.items():
            text = text.replace(err, corr)
        text = re.sub(r'\b0\b(?=\s*[a-zA-Z])', 'O', text)
        text = re.sub(r'\bl\b(?=\s*\d)', '1', text)
        return text

    def normalize_spacing(self, text):
        return re.sub(r'\s+', ' ', text).strip()

    def clean_text(self, text):
        text = self.remove_noise_patterns(text)
        text = self.correct_ocr_errors(text)
        text = self.normalize_spacing(text)
        return text


# =========================
# Table helpers
# =========================
def forward_fill(row):
    last = ""
    filled = []
    for val in row:
        if val and str(val).strip():
            last = str(val).strip()
        filled.append(last)
    return filled

def merge_headers(header_rows):
    merged = []
    num_cols = max(len(r) for r in header_rows)
    for i in range(num_cols):
        parts = [str(r[i]).strip() for r in header_rows if i < len(r) and r[i] and str(r[i]).strip()]
        merged.append(" ".join(parts).strip())
    return merged

def is_header_row(row):
    if not row:
        return False
    filled_count = sum(1 for cell in row if cell and str(cell).strip())
    return filled_count >= len(row) / 2


async def fetch_document(url: str, timeout: int = ASYNC_TIMEOUT):
    safe_url = url.split('?')[0]
    logger.info(f"Starting download of PDF from {safe_url}")
    file_ext = safe_url.split('.')[-1].lower()
    logger.debug(f"Detected file extension: {file_ext}")

    # Save to project-level temp folder
    temp_dir = Path(__file__).resolve().parent.parent / "temp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    file_path = temp_dir / f"{int(time.time())}.{file_ext}"
    temp_path = file_path

    headers = {'User-Agent': 'HackRx-RAG-System/1.0'}

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=timeout, headers=headers) as response:
                if response.status != 200:
                    logger.error(f"Download failed: HTTP {response.status}")
                    return None
                
                with open(temp_path, 'wb') as f:
                    async for chunk in response.content.iter_chunked(4096):
                        f.write(chunk)

        logger.info(f"Downloaded {file_ext.upper()} to {temp_path}")
        return temp_path, file_ext
    
    except asyncio.TimeoutError:
        logger.error(f"Timeout after {timeout}s")
    except Exception as e:
        logger.error(f"Error: {e}")
    
    try:
        temp_path.unlink()
    except Exception:
        pass
    
    return None


async def parse_pdf(path):
    documents = await parser.aload_data(path)
    return documents
def parse_docx(path):
    cleaner = UniversalTextCleaner()
    doc = docx.Document(path)

    # Text
    raw_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    cleaned_text = cleaner.clean_text(raw_text)

    # Tables
    table_rows = []
    for table in doc.tables:
        headers = forward_fill([cell.text.strip() for cell in table.rows[0].cells])
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            if len(row_data) != len(headers):
                continue
            pairs = [f"{h}: {c}" for h, c in zip(headers, row_data)
                     if h and c and str(c).strip().lower() != "none"]
            if pairs:
                table_rows.append(", ".join(pairs))

    return cleaned_text, table_rows

def parse_email(path):
    cleaner = UniversalTextCleaner()
    with open(path, 'rb') as f:
        msg = email.message_from_binary_file(f, policy=policy.default)

    text_parts = []
    table_rows = []

    for part in msg.walk():
        content_type = part.get_content_type()
        if content_type == "text/plain":
            text_parts.append(part.get_content())
        elif content_type == "text/html":
            html_content = part.get_content()
            soup = BeautifulSoup(html_content, "html.parser")
            text_parts.append(soup.get_text())

            # Extract HTML tables
            for html_table in soup.find_all("table"):
                rows = []
                for tr in html_table.find_all("tr"):
                    cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                    rows.append(cells)
                if rows and len(rows) > 1:
                    headers = forward_fill(rows[0])
                    for row in rows[1:]:
                        if len(row) != len(headers):
                            continue
                        pairs = [f"{h}: {c}" for h, c in zip(headers, row)
                                 if h and c and str(c).strip().lower() != "none"]
                        if pairs:
                            table_rows.append(", ".join(pairs))

    cleaned_text = cleaner.clean_text("\n".join(text_parts))
    return cleaned_text, table_rows

async def save_file(content):
    temp_dir = Path(__file__).resolve().parent.parent / "temp"
    temp_dir.mkdir(exist_ok=True)
    
    filename = f"pdf_to_text.txt"

    file_path = temp_dir / filename

    # Save content
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(str(content))

    logger.info(f"Saved parsed content to {file_path}")
    return str(file_path) 

async def document_downloader(url: str):
    """
    Download a document (PDF, DOCX, or email) and extract its text.

    Args:
        url: URL of the document

    Returns:
        Extracted text as string, or None if processing fails
    """
    safe_url = url
    if '%' not in url and ('&' in url or '+' in url or '=' in url):
        # Only encode the query part, not the entire URL
        parts = url.split('?', 1)
        if len(parts) > 1:
            base_url, query = parts
            # Don't re-encode percent signs
            safe_url = f"{base_url}?{query}"
    try:
        result = await fetch_document(url)
        if not result:
            logger.error("Failed to download document")
            return None

        doc_path, file_ext = result

        try:
            if file_ext == "pdf":
                final_output =  await parse_pdf(doc_path)


            elif file_ext in ["docx", "doc"]:
                text,table = parse_docx(doc_path)
                final_output = text.strip()
                if table:
                    final_output += "; " + "; ".join(table)


            elif file_ext in ["eml", "msg"]:
                logger.debug(f"Processing email file: {doc_path}")
                text,table = parse_email(doc_path)
                final_output = text.strip()
                if table:
                    final_output += "; " + "; ".join(table)
                    
            else:
                logger.error(f"Unsupported file type: {file_ext}")
                return None
            saved_path = await save_file(final_output)
            logger.info(f"Document saved at: {saved_path}")
            return saved_path
        finally:
            try:
                Path(doc_path).unlink(missing_ok=True)
            except Exception as e:
                logger.warning(f"Failed to delete temp file: {e}")

    except Exception as e:
        logger.error(f"Error processing document: {e}")
        return None
