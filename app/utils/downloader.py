import os
import re
import asyncio
import tempfile
from pathlib import Path
from collections import Counter
import logging
import aiohttp
import fitz  # PyMuPDF
import pdfplumber
import docx
import email
from email import policy
from bs4 import BeautifulSoup

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =========================
# Text Cleaner Utility
# =========================
class TextSanitizer:
    def __init__(self):
        self.patterns_to_remove = [
            r'UIN[:\-\s]*[A-Z0-9]+',
            r'Reg\.?\s*No\.?\s*[:\-\s]*\d+',
            r'CIN\s*[:\-\s]*[A-Z0-9]+',
            r'IRDAI\s*Regn?\.?\s*No\.?\s*[:\-\s]*\d+',
            r'Page\s+\d+\s+of\s+\d+',
            r'^\d+\s*$',
            r'www\.[a-zA-Z0-9\-\.]+\.com',
            r'E-mail:\s*[a-zA-Z0-9\-\.@]+',
            r'Call at:\s*.*?\(Toll Free.*?\)',
            r'For more details.*?Toll Free.*?\)',
            r'Regd\.?\s*&\s*Head Office:.*?-\s*\d{6}',
            r'Plot no\..*?-\s*\d{6}',
            r'^[A-Z\s]{10,}\s*$',
        ]
        self.ocr_fixes = {
            ' . ': '. ', ' , ': ', ', ' ; ': '; ', ' : ': ': ',
            '( ': '(', ' )': ')', ' / ': '/',
            'O ': '0 ', 'l ': '1 ', '  ': ' ',
        }

    def _identify_repeated_lines(self, lines, threshold=3):
        line_counts = Counter(line.strip() for line in lines if len(line.strip()) > 10)
        return {line for line, count in line_counts.items() if count >= threshold}

    def _remove_patterns(self, text):
        lines = text.split("\n")
        repeated = self._identify_repeated_lines(lines)
        cleaned = []
        for line in lines:
            line = line.strip()
            if not line or line in repeated:
                continue
            if any(re.search(p, line, re.IGNORECASE) for p in self.patterns_to_remove):
                continue
            cleaned.append(line)
        return " ".join(cleaned)

    def _fix_common_ocr_errors(self, text):
        for wrong, right in self.ocr_fixes.items():
            text = text.replace(wrong, right)
        text = re.sub(r'\b0\b(?=\s*[a-zA-Z])', 'O', text)
        text = re.sub(r'\bl\b(?=\s*\d)', '1', text)
        return text

    def _normalize_whitespace(self, text):
        return re.sub(r'\s+', ' ', text).strip()

    def clean(self, text):
        return self._normalize_whitespace(
            self._fix_common_ocr_errors(
                self._remove_patterns(text)))


# =========================
# Table Utilities
# =========================
def forward_fill_row(row):
    last_value = ""
    result = []
    for cell in row:
        if cell and str(cell).strip():
            last_value = str(cell).strip()
        result.append(last_value)
    return result

def merge_table_headers(header_rows):
    max_columns = max(len(r) for r in header_rows)
    return [" ".join(str(r[i]).strip() for r in header_rows if i < len(r) and r[i]).strip() for i in range(max_columns)]

def is_likely_header(row):
    filled = sum(1 for cell in row if cell and str(cell).strip())
    return filled >= len(row) / 2


# =========================
# Format-Specific Parsers
# =========================
def parse_pdf(path):
    sanitizer = TextSanitizer()
    with fitz.open(path) as doc:
        raw_text = "\n\n".join(page.get_text("text") for page in doc)
    cleaned_text = sanitizer.clean(raw_text)

    rows = []
    seen = set()
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                if not table or len(table) < 2:
                    continue
                headers = []
                start_row = 0
                for i in range(min(3, len(table))):
                    if is_likely_header(table[i]):
                        headers.append(table[i])
                        start_row = i + 1
                    else:
                        break
                if not headers:
                    headers = [table[0]]
                    start_row = 1
                merged = forward_fill_row(merge_table_headers(headers))
                for row in table[start_row:]:
                    row_data = [str(cell).replace("\n", " ").strip() if cell else "" for cell in forward_fill_row(row)]
                    if len(row_data) != len(merged):
                        continue
                    entry = ", ".join(f"{h}: {v}" for h, v in zip(merged, row_data) if h and v)
                    if entry and entry not in seen:
                        seen.add(entry)
                        rows.append(entry)
    return cleaned_text, rows

def parse_docx(path):
    sanitizer = TextSanitizer()
    doc = docx.Document(path)
    raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    cleaned_text = sanitizer.clean(raw_text)

    rows = []
    for table in doc.tables:
        headers = forward_fill_row([c.text.strip() for c in table.rows[0].cells])
        for row in table.rows[1:]:
            row_data = [c.text.strip() for c in row.cells]
            if len(row_data) != len(headers):
                continue
            pairs = [f"{h}: {v}" for h, v in zip(headers, row_data) if h and v]
            if pairs:
                rows.append(", ".join(pairs))
    return cleaned_text, rows

def parse_eml(path):
    sanitizer = TextSanitizer()
    with open(path, 'rb') as f:
        msg = email.message_from_binary_file(f, policy=policy.default)

    text_parts, rows = [], []
    for part in msg.walk():
        if part.get_content_type() == "text/plain":
            text_parts.append(part.get_content())
        elif part.get_content_type() == "text/html":
            soup = BeautifulSoup(part.get_content(), "html.parser")
            text_parts.append(soup.get_text())
            for table in soup.find_all("table"):
                html_rows = [[td.get_text(strip=True) for td in tr.find_all(["td", "th"])] for tr in table.find_all("tr")]
                if html_rows and len(html_rows) > 1:
                    headers = forward_fill_row(html_rows[0])
                    for row in html_rows[1:]:
                        if len(row) != len(headers):
                            continue
                        entry = ", ".join(f"{h}: {c}" for h, c in zip(headers, row) if h and c)
                        if entry:
                            rows.append(entry)
    cleaned_text = sanitizer.clean("\n".join(text_parts))
    return cleaned_text, rows


# =========================
# File Handler
# =========================
def parse_local_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".eml":
        return parse_eml(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def save_output(cleaned_text, table_data, output_path):
    final_output = cleaned_text.strip()
    if table_data:
        final_output += "; " + "; ".join(table_data)
    final_output = final_output.replace("\n", " ")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(final_output)
    logger.info(f"Saved output to {output_path} with {len(table_data)} table rows")
    return output_path


async def fetch_document(url: str, output_file: str = None):
    if not output_file:
        output_file = "extracted_content.txt"

    try:
        logger.info(f"Fetching document from {url}")
        suffix = '.' + url.split('?')[0].split('.')[-1].lower()
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        path = Path(temp_file.name)

        headers = {'User-Agent': 'DocFetcher/1.0'}
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=30, headers=headers) as response:
                if response.status != 200:
                    logger.error(f"Failed to fetch file: HTTP {response.status}")
                    return None
                with open(path, 'wb') as f:
                    async for chunk in response.content.iter_chunked(4096):
                        f.write(chunk)

        text, tables = parse_local_file(str(path))
        output = save_output(text, tables, output_file)
        return output

    except Exception as e:
        logger.error(f"Error: {e}")
    finally:
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass
    return None


def parse_file(file_path: str, output_file: str = None):
    if not output_file:
        output_file = "extracted_content.txt"

    if not os.path.exists(file_path):
        logger.error(f"File does not exist: {file_path}")
        return None

    try:
        text, tables = parse_local_file(file_path)
        return save_output(text, tables, output_file)
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return None
